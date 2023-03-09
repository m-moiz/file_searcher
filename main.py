import os
import openpyxl
import docx


def search():
    search_term = input("Enter search term: ")
    directory = input("Enter directory to search: ")

    # Identify the files to search
    files_to_search = []
    for filename in os.listdir(directory):
        if filename.endswith(".xlsx") or filename.endswith(".docx") or filename.endswith(".txt"):
            files_to_search.append(os.path.join(directory, filename))

    # Search for the term in each file
    results = []
    for file_path in files_to_search:
        if file_path.endswith(".xlsx"):
            row_values = search_in_xsl_files(file_path,search_term)
            results.append((file_path, sheet_name, row_values))

        elif file_path.endswith(".docx"):
            result = search_in_doc_files(file_path, search_term)
            results.append((file_path, 0, result))

    # Display the results
    if results:
        print("Results:")
        for result in results:
            file_path, sheet_name, row = result
            if file_path.endswith(".xlsx"):
                for cell in row:
                    print(f"{file_path}: {sheet_name} - {cell}")
            elif file_path.endswith(".docx"):
                file_path, i, text = result
                print(f"{file_path}: {i}, {text}\n")
    else:
        print("No results found.")

def search_in_xsl_files(file_path, search_term):
    # Read Excel file
    workbook = openpyxl.load_workbook(file_path)
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        for i, row in enumerate(sheet.rows):
            first_cell_value = row[0].value
            if first_cell_value and type(first_cell_value) == str and search_term.lower() in first_cell_value.lower():
                row_values = sheet.iter_rows(min_row=i + 1, max_row=i + 1, values_only=True)


def search_in_doc_files(file_path, search_term):
    doc = docx.Document(file_path)
    index = 0
    result = []
    for i, para in enumerate(doc.paragraphs):
        words = para.text.split()
        if search_term.lower() in para.text.lower():
            index = i
            result = [para.text]
            # Keep appending subsequent lines until a blank line is encountered
            while True:
                para = doc.paragraphs[i + 1]
                if para.text.strip() == '':
                    break
                result.append(para.text)
                i += 1
    return result

# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    search()

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
